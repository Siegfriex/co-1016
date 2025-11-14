#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VID Markdown to DOCX Converter
Converts VID.md to a formatted DOCX document with visual design elements.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

def add_color_cell(cell, hex_color):
    """Add a colored cell background"""
    try:
        shading = OxmlElement('w:shd')
        hex_clean = hex_color.replace('#', '').upper()
        shading.set(qn('w:fill'), hex_clean)
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    except Exception as e:
        print(f"Warning: Could not apply color {hex_color}: {e}")

def rgb_from_hex(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.replace('#', '')
    if len(hex_color) == 6:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return (0, 0, 0)

def parse_markdown_table(line):
    """Parse markdown table line"""
    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
    return cells

def is_table_separator(line):
    """Check if line is a markdown table separator"""
    return re.match(r'^\|[\s\-:]+\|', line)

def set_korean_font(run, font_name='맑은 고딕', size=11):
    """Set Korean font for a text run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    # Set East Asian font for Korean characters
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    except Exception:
        pass  # Fallback if font setting fails

def convert_vid_to_docx(md_file, docx_file):
    """Convert VID.md to DOCX"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    # Set East Asian font
    try:
        rPr = font.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:ascii'), '맑은 고딕')
        rFonts.set(qn('w:hAnsi'), '맑은 고딕')
    except Exception:
        pass
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    table_data = []
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_para = doc.add_paragraph()
                code_run = code_para.add_run('\n'.join(code_block_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0, 100, 0)
                # Keep Consolas for code, but ensure Korean characters are handled
                try:
                    rPr = code_run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                except Exception:
                    pass
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_language = line.replace('```', '').strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                para = doc.add_heading(text, level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 24)
                    para.runs[0].bold = True
            elif level == 2:
                para = doc.add_heading(text, level=2)
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 18)
                    para.runs[0].bold = True
            elif level == 3:
                para = doc.add_heading(text, level=3)
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 14)
                    para.runs[0].bold = True
            elif level == 4:
                para = doc.add_heading(text, level=4)
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 12)
                    para.runs[0].bold = True
            else:
                para = doc.add_paragraph(text)
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 11)
                    para.runs[0].bold = True
            
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not line.strip().startswith('|'):
            # Check if next line is table separator
            if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                # Start of table
                header = parse_markdown_table(line)
                i += 1  # Skip separator
                table_data = [header]
                
                # Collect table rows
                while i < len(lines):
                    next_line = lines[i].rstrip()
                    if '|' in next_line and not is_table_separator(next_line):
                        table_data.append(parse_markdown_table(next_line))
                        i += 1
                    else:
                        break
                
                # Create table
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx >= len(table.rows[row_idx].cells):
                                continue
                            cell = table.rows[row_idx].cells[col_idx]
                            
                            # Check for color codes and apply background BEFORE setting text
                            color_match = re.search(r'#([0-9A-Fa-f]{6})', cell_data)
                            if color_match:
                                hex_color = '#' + color_match.group(1)
                                add_color_cell(cell, hex_color)
                            
                            # Set cell text (remove color codes for display)
                            display_text = re.sub(r'`#([0-9A-Fa-f]{6})`', r'#\1', cell_data)
                            cell.text = display_text
                            
                            # Style header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        set_korean_font(run, '맑은 고딕', 10)
                                        run.bold = True
                                        # Make text white if background is dark
                                        if color_match:
                                            rgb = rgb_from_hex(hex_color)
                                            if sum(rgb) < 400:  # Dark background
                                                run.font.color.rgb = RGBColor(255, 255, 255)
                            
                            # Style cells
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    set_korean_font(run, '맑은 고딕', 9)
                                    # Make text readable on colored backgrounds
                                    if color_match and row_idx > 0:
                                        rgb = rgb_from_hex(hex_color)
                                        if sum(rgb) < 400:  # Dark background
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                        else:
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    table_data = []
                continue
        
        # Handle bold text
        if '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part.strip('*'))
                    set_korean_font(run, '맑은 고딕', 11)
                    run.bold = True
                else:
                    run = para.add_run(part)
                    set_korean_font(run, '맑은 고딕', 11)
            i += 1
            continue
        
        # Handle list items
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            if para.runs:
                set_korean_font(para.runs[0], '맑은 고딕', 11)
            i += 1
            continue
        
        # Handle numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            para = doc.add_paragraph(text, style='List Number')
            if para.runs:
                set_korean_font(para.runs[0], '맑은 고딕', 11)
            i += 1
            continue
        
        # Handle ASCII diagrams (convert to formatted text)
        if line.strip() and ('┌' in line or '│' in line or '└' in line or '─' in line):
            # ASCII diagram - use monospace font
            para = doc.add_paragraph(line)
            para.runs[0].font.name = 'Consolas'
            para.runs[0].font.size = Pt(9)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            # Check for inline code
            if '`' in line:
                para = doc.add_paragraph()
                parts = re.split(r'(`[^`]+`)', line)
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = para.add_run(part.strip('`'))
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 100, 0)
                    else:
                        run = para.add_run(part)
                        set_korean_font(run, '맑은 고딕', 11)
            else:
                para = doc.add_paragraph(line)
                if para.runs:
                    set_korean_font(para.runs[0], '맑은 고딕', 11)
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    md_file = 'docs/design/VID.md'
    docx_file = 'docs/design/VID.docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    
    convert_vid_to_docx(md_file, docx_file)

